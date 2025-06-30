"""
Report Generator Module
======================

This module generates professional PDF and Word reports from cryptocurrency
address extraction results.

Author: Crypto Extractor Tool
Date: 2025-01-09
Version: 1.0.0
"""

import os
import logging
from typing import List, Dict, Optional
from datetime import datetime
from collections import Counter

# For PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
from reportlab.platypus import KeepTogether
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.charts.barcharts import VerticalBarChart

# For Word generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from extractor import ExtractedAddress


class ReportGenerator:
    """
    Generates professional reports in PDF and Word formats.
    """
    
    def __init__(self):
        """Initialize the report generator."""
        self.logger = logging.getLogger(__name__)
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom paragraph styles for the report."""
        # Title style
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor('#1a237e'),
            alignment=TA_CENTER
        )
        
        # Heading style
        self.heading_style = ParagraphStyle(
            'CustomHeading',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.HexColor('#1976d2'),
            spaceBefore=12
        )
        
        # Subheading style
        self.subheading_style = ParagraphStyle(
            'CustomSubheading',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            textColor=colors.HexColor('#424242')
        )
        
        # Normal text style
        self.normal_style = ParagraphStyle(
            'CustomNormal',
            parent=self.styles['Normal'],
            fontSize=10,
            leading=14
        )
    
    def generate_pdf_report(self, addresses: List[ExtractedAddress], output_path: str, 
                           metadata: Optional[Dict] = None) -> str:
        """
        Generate a PDF report from extraction results.
        
        Args:
            addresses: List of extracted addresses
            output_path: Path for the output PDF file
            metadata: Optional metadata (extraction date, files processed, etc.)
        
        Returns:
            Path to the generated PDF file
        """
        try:
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Build the story (content)
            story = []
            
            # Title page
            story.extend(self._create_pdf_title_page(metadata))
            story.append(PageBreak())
            
            # Executive summary
            story.extend(self._create_pdf_summary(addresses, metadata))
            story.append(PageBreak())
            
            # Statistics section
            story.extend(self._create_pdf_statistics(addresses))
            story.append(PageBreak())
            
            # Detailed findings by cryptocurrency
            story.extend(self._create_pdf_detailed_findings(addresses))
            
            # Build the PDF
            doc.build(story)
            
            self.logger.info(f"PDF report generated: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to generate PDF report: {str(e)}")
            raise
    
    def _create_pdf_title_page(self, metadata: Optional[Dict]) -> List:
        """Create the title page for the PDF report."""
        elements = []
        
        # Title
        elements.append(Spacer(1, 2*inch))
        elements.append(Paragraph("Cryptocurrency Address Extraction Report", self.title_style))
        elements.append(Spacer(1, 0.5*inch))
        
        # Subtitle with date
        date_str = datetime.now().strftime("%B %d, %Y")
        elements.append(Paragraph(f"Generated on {date_str}", self.subheading_style))
        
        # Metadata if provided
        if metadata:
            elements.append(Spacer(1, inch))
            
            if 'organization' in metadata:
                elements.append(Paragraph(f"Organization: {metadata['organization']}", self.normal_style))
            
            if 'analyst' in metadata:
                elements.append(Paragraph(f"Analyst: {metadata['analyst']}", self.normal_style))
            
            if 'case_number' in metadata:
                elements.append(Paragraph(f"Case Number: {metadata['case_number']}", self.normal_style))
        
        return elements
    
    def _create_pdf_summary(self, addresses: List[ExtractedAddress], metadata: Optional[Dict]) -> List:
        """Create the executive summary section."""
        elements = []
        
        elements.append(Paragraph("Executive Summary", self.heading_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Calculate statistics
        total_addresses = len(addresses)
        unique_addresses = len(set((addr.address, addr.crypto_type) for addr in addresses))
        duplicate_count = total_addresses - unique_addresses
        
        # Get crypto breakdown
        crypto_counts = Counter(addr.crypto_name for addr in addresses)
        
        # Summary text
        summary_text = f"""
        This report presents the results of cryptocurrency address extraction performed on 
        {metadata.get('files_processed', 'multiple')} CSV files. The extraction process identified 
        a total of {total_addresses} cryptocurrency addresses, of which {unique_addresses} are unique.
        
        Key Findings:
        • Total addresses found: {total_addresses}
        • Unique addresses: {unique_addresses}
        • Duplicate addresses: {duplicate_count}
        • Cryptocurrency types identified: {len(crypto_counts)}
        
        The most common cryptocurrency was {crypto_counts.most_common(1)[0][0]} with 
        {crypto_counts.most_common(1)[0][1]} addresses found.
        """
        
        elements.append(Paragraph(summary_text, self.normal_style))
        
        # Add a summary table
        elements.append(Spacer(1, 0.3*inch))
        elements.append(Paragraph("Summary by Cryptocurrency", self.subheading_style))
        elements.append(Spacer(1, 0.1*inch))
        
        # Create summary table
        table_data = [['Cryptocurrency', 'Count', 'Percentage']]
        for crypto, count in crypto_counts.most_common():
            percentage = f"{(count/total_addresses*100):.1f}%"
            table_data.append([crypto, str(count), percentage])
        
        summary_table = Table(table_data, colWidths=[3*inch, 1.5*inch, 1.5*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1976d2')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
        ]))
        
        elements.append(summary_table)
        
        return elements
    
    def _create_pdf_statistics(self, addresses: List[ExtractedAddress]) -> List:
        """Create statistics section with charts."""
        elements = []
        
        elements.append(Paragraph("Statistical Analysis", self.heading_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Confidence distribution
        elements.append(Paragraph("Confidence Score Distribution", self.subheading_style))
        
        # Group by confidence ranges
        confidence_ranges = {
            'High (90-100%)': 0,
            'Medium (70-89%)': 0,
            'Low (Below 70%)': 0
        }
        
        for addr in addresses:
            if addr.confidence >= 90:
                confidence_ranges['High (90-100%)'] += 1
            elif addr.confidence >= 70:
                confidence_ranges['Medium (70-89%)'] += 1
            else:
                confidence_ranges['Low (Below 70%)'] += 1
        
        # Create confidence table
        conf_data = [['Confidence Level', 'Count', 'Percentage']]
        total = len(addresses)
        for level, count in confidence_ranges.items():
            percentage = f"{(count/total*100):.1f}%" if total > 0 else "0%"
            conf_data.append([level, str(count), percentage])
        
        conf_table = Table(conf_data, colWidths=[3*inch, 1.5*inch, 1.5*inch])
        conf_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4caf50')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
        ]))
        
        elements.append(conf_table)
        
        # Duplicate analysis
        elements.append(Spacer(1, 0.3*inch))
        elements.append(Paragraph("Duplicate Address Analysis", self.subheading_style))
        
        # Find most duplicated addresses
        address_counts = Counter(addr.address for addr in addresses)
        most_common = address_counts.most_common(10)
        
        if most_common:
            dup_data = [['Address', 'Occurrences', 'Type']]
            for address, count in most_common:
                if count > 1:
                    # Find the crypto type
                    crypto_type = next((a.crypto_name for a in addresses if a.address == address), 'Unknown')
                    # Truncate address for display
                    display_addr = f"{address[:20]}...{address[-10:]}" if len(address) > 35 else address
                    dup_data.append([display_addr, str(count), crypto_type])
            
            if len(dup_data) > 1:
                dup_table = Table(dup_data, colWidths=[3.5*inch, 1.25*inch, 1.25*inch])
                dup_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#ff9800')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
                    ('FONTSIZE', (0, 1), (0, -1), 9),  # Smaller font for addresses
                ]))
                elements.append(dup_table)
            else:
                elements.append(Paragraph("No duplicate addresses found.", self.normal_style))
        
        return elements
    
    def _create_pdf_detailed_findings(self, addresses: List[ExtractedAddress]) -> List:
        """Create detailed findings section organized by cryptocurrency."""
        elements = []
        
        elements.append(Paragraph("Detailed Findings", self.heading_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Group addresses by cryptocurrency
        from collections import defaultdict
        crypto_groups = defaultdict(list)
        for addr in addresses:
            crypto_groups[addr.crypto_name].append(addr)
        
        # Create section for each cryptocurrency
        for crypto_name in sorted(crypto_groups.keys()):
            crypto_addresses = crypto_groups[crypto_name]
            
            # Cryptocurrency header
            elements.append(Paragraph(f"{crypto_name} Addresses", self.subheading_style))
            elements.append(Paragraph(f"Total found: {len(crypto_addresses)}", self.normal_style))
            elements.append(Spacer(1, 0.1*inch))
            
            # Sample addresses table (first 10)
            sample_data = [['Address', 'File', 'Confidence', 'Duplicate']]
            
            for addr in crypto_addresses[:10]:
                display_addr = f"{addr.address[:25]}..." if len(addr.address) > 30 else addr.address
                display_file = addr.filename if len(addr.filename) <= 20 else f"...{addr.filename[-17:]}"
                is_dup = "Yes" if addr.is_duplicate else "No"
                sample_data.append([
                    display_addr,
                    display_file,
                    f"{addr.confidence:.0f}%",
                    is_dup
                ])
            
            sample_table = Table(sample_data, colWidths=[3*inch, 2*inch, 0.75*inch, 0.75*inch])
            sample_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#757575')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (2, 0), (3, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#fafafa')]),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
            ]))
            
            elements.append(KeepTogether(sample_table))
            
            if len(crypto_addresses) > 10:
                elements.append(Paragraph(
                    f"... and {len(crypto_addresses) - 10} more addresses", 
                    self.normal_style
                ))
            
            elements.append(Spacer(1, 0.3*inch))
        
        return elements
    
    def generate_word_report(self, addresses: List[ExtractedAddress], output_path: str,
                            metadata: Optional[Dict] = None) -> str:
        """
        Generate a Word document report from extraction results.
        
        Args:
            addresses: List of extracted addresses
            output_path: Path for the output Word file
            metadata: Optional metadata
        
        Returns:
            Path to the generated Word file
        """
        try:
            doc = Document()
            
            # Set up document properties
            self._setup_word_document(doc)
            
            # Title page
            self._create_word_title_page(doc, metadata)
            
            # Executive summary
            self._create_word_summary(doc, addresses, metadata)
            
            # Statistics
            self._create_word_statistics(doc, addresses)
            
            # Detailed findings
            self._create_word_detailed_findings(doc, addresses)
            
            # Save the document
            doc.save(output_path)
            
            self.logger.info(f"Word report generated: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"Failed to generate Word report: {str(e)}")
            raise
    
    def _setup_word_document(self, doc: Document):
        """Set up Word document styles and properties."""
        # Modify default styles
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)
        
        doc.styles['Heading 1'].font.color.rgb = RGBColor(26, 35, 126)  # Dark blue
        doc.styles['Heading 1'].font.size = Pt(18)
        
        doc.styles['Heading 2'].font.color.rgb = RGBColor(25, 118, 210)  # Blue
        doc.styles['Heading 2'].font.size = Pt(14)
    
    def _create_word_title_page(self, doc: Document, metadata: Optional[Dict]):
        """Create title page for Word document."""
        # Title
        title = doc.add_heading('Cryptocurrency Address Extraction Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(14)
        
        # Metadata
        if metadata:
            doc.add_paragraph()
            doc.add_paragraph()
            
            if 'organization' in metadata:
                org_para = doc.add_paragraph()
                org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                org_para.add_run(f"Organization: {metadata['organization']}")
            
            if 'analyst' in metadata:
                analyst_para = doc.add_paragraph()
                analyst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                analyst_para.add_run(f"Analyst: {metadata['analyst']}")
        
        # Page break
        doc.add_page_break()
    
    def _create_word_summary(self, doc: Document, addresses: List[ExtractedAddress], 
                            metadata: Optional[Dict]):
        """Create executive summary in Word document."""
        doc.add_heading('Executive Summary', 1)
        
        # Statistics
        total_addresses = len(addresses)
        unique_addresses = len(set((addr.address, addr.crypto_type) for addr in addresses))
        duplicate_count = total_addresses - unique_addresses
        crypto_counts = Counter(addr.crypto_name for addr in addresses)
        
        # Summary paragraph
        summary_text = (
            f"This report presents the results of cryptocurrency address extraction performed on "
            f"{metadata.get('files_processed', 'multiple')} CSV files. The extraction process identified "
            f"a total of {total_addresses} cryptocurrency addresses, of which {unique_addresses} are unique."
        )
        doc.add_paragraph(summary_text)
        
        # Key findings list
        doc.add_paragraph("Key Findings:")
        doc.add_paragraph(f"• Total addresses found: {total_addresses}", style='List Bullet')
        doc.add_paragraph(f"• Unique addresses: {unique_addresses}", style='List Bullet')
        doc.add_paragraph(f"• Duplicate addresses: {duplicate_count}", style='List Bullet')
        doc.add_paragraph(f"• Cryptocurrency types identified: {len(crypto_counts)}", style='List Bullet')
        
        doc.add_paragraph()
        
        # Summary table
        doc.add_heading('Summary by Cryptocurrency', 2)
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Cryptocurrency'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage'
        
        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data rows
        for crypto, count in crypto_counts.most_common():
            row_cells = table.add_row().cells
            row_cells[0].text = crypto
            row_cells[1].text = str(count)
            row_cells[2].text = f"{(count/total_addresses*100):.1f}%"
        
        doc.add_page_break()
    
    def _create_word_statistics(self, doc: Document, addresses: List[ExtractedAddress]):
        """Create statistics section in Word document."""
        doc.add_heading('Statistical Analysis', 1)
        
        # Confidence distribution
        doc.add_heading('Confidence Score Distribution', 2)
        
        confidence_ranges = {
            'High (90-100%)': sum(1 for a in addresses if a.confidence >= 90),
            'Medium (70-89%)': sum(1 for a in addresses if 70 <= a.confidence < 90),
            'Low (Below 70%)': sum(1 for a in addresses if a.confidence < 70)
        }
        
        # Create confidence table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Confidence Level'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage'
        
        total = len(addresses)
        for level, count in confidence_ranges.items():
            row_cells = table.add_row().cells
            row_cells[0].text = level
            row_cells[1].text = str(count)
            row_cells[2].text = f"{(count/total*100):.1f}%" if total > 0 else "0%"
        
        doc.add_paragraph()
        
        # Duplicate analysis
        doc.add_heading('Top Duplicate Addresses', 2)
        
        address_counts = Counter(addr.address for addr in addresses)
        most_common = [(addr, count) for addr, count in address_counts.most_common(10) if count > 1]
        
        if most_common:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Address'
            hdr_cells[1].text = 'Occurrences'
            hdr_cells[2].text = 'Type'
            
            for address, count in most_common[:5]:  # Top 5 duplicates
                crypto_type = next((a.crypto_name for a in addresses if a.address == address), 'Unknown')
                display_addr = f"{address[:20]}...{address[-10:]}" if len(address) > 35 else address
                
                row_cells = table.add_row().cells
                row_cells[0].text = display_addr
                row_cells[1].text = str(count)
                row_cells[2].text = crypto_type
        else:
            doc.add_paragraph("No duplicate addresses found.")
        
        doc.add_page_break()
    
    def _create_word_detailed_findings(self, doc: Document, addresses: List[ExtractedAddress]):
        """Create detailed findings section in Word document."""
        doc.add_heading('Detailed Findings', 1)
        
        # Group by cryptocurrency
        from collections import defaultdict
        crypto_groups = defaultdict(list)
        for addr in addresses:
            crypto_groups[addr.crypto_name].append(addr)
        
        for crypto_name in sorted(crypto_groups.keys()):
            crypto_addresses = crypto_groups[crypto_name]
            
            doc.add_heading(f"{crypto_name} Addresses", 2)
            doc.add_paragraph(f"Total found: {len(crypto_addresses)}")
            
            # Sample addresses table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light List Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Address'
            hdr_cells[1].text = 'File'
            hdr_cells[2].text = 'Confidence'
            hdr_cells[3].text = 'Duplicate'
            
            # Show first 5 addresses
            for addr in crypto_addresses[:5]:
                display_addr = f"{addr.address[:25]}..." if len(addr.address) > 30 else addr.address
                display_file = addr.filename if len(addr.filename) <= 20 else f"...{addr.filename[-17:]}"
                
                row_cells = table.add_row().cells
                row_cells[0].text = display_addr
                row_cells[1].text = display_file
                row_cells[2].text = f"{addr.confidence:.0f}%"
                row_cells[3].text = "Yes" if addr.is_duplicate else "No"
            
            if len(crypto_addresses) > 5:
                doc.add_paragraph(f"... and {len(crypto_addresses) - 5} more addresses")
            
            doc.add_paragraph()